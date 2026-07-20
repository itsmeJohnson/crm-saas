"""AI Document Intelligence — OCR, PDF/DOCX/XLSX/image/text parsing, document
classification, structured data extraction (invoice / contract / identity /
resume), table extraction, image understanding, AI summaries and semantic
document search.

The pipeline is deterministic-first: parsing, classification, extraction and
tables need no AI provider. OCR is pluggable (pytesseract + tesseract binary
when installed; capabilities are reported honestly when not). AI summaries go
through AIGatewayService — never a provider directly. Search reuses the
Knowledge Base's hash_embed_v1 embedding pipeline.
"""
import csv
import io
import re
import uuid
from datetime import datetime, timezone

from fastapi import HTTPException, status
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.user import User
from app.models.document_intelligence import DIDocument
from app.services.audit_service import AuditService
from app.services.knowledge_base_service import embed_text, cosine, _tokenize, EMBED_MODEL

MANAGER_ROLES = ("SuperAdmin", "OrgAdmin", "Manager")
DOC_TYPES = ("invoice", "contract", "identity", "resume", "receipt", "report", "letter", "other")
TEXT_CAP = 100_000
IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif", ".webp")


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _aware(dt):
    if dt is None:
        return None
    return dt if dt.tzinfo else dt.replace(tzinfo=timezone.utc)


# ---------- capability probes ----------
_OCR_STATE: dict = {}


def ocr_available() -> bool:
    if "ok" not in _OCR_STATE:
        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            _OCR_STATE["ok"] = True
        except Exception:
            _OCR_STATE["ok"] = False
    return _OCR_STATE["ok"]


def _lib_available(module: str) -> bool:
    try:
        __import__(module)
        return True
    except ImportError:
        return False


def capabilities() -> dict:
    return {
        "pdf": _lib_available("pypdf"),
        "docx": _lib_available("docx"),
        "xlsx": _lib_available("openpyxl"),
        "images": _lib_available("PIL"),
        "ocr": ocr_available(),
        "text_formats": [".txt", ".csv", ".md", ".json", ".log", ".html"],
        "image_formats": list(IMAGE_EXTS),
        "embedding_model": EMBED_MODEL,
    }


# ---------- parsing (OCR / PDF / DOCX / XLSX / image / text) ----------
def _parse_pdf(content: bytes) -> dict:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    pages = [(p.extract_text() or "") for p in reader.pages]
    text = "\n\n".join(pages).strip()
    out = {"text": text, "page_count": len(pages), "parser": "pypdf"}
    if len(pages) > 0 and len(text) < 20:
        # no text layer -> scanned PDF; page rasterization is not available,
        # so OCR cannot run on it even when tesseract is installed
        out["needs_ocr"] = True
    return out


def _parse_docx(content: bytes) -> dict:
    import docx
    d = docx.Document(io.BytesIO(content))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    tables = []
    for t in d.tables:
        rows = [[(c.text or "").strip() for c in row.cells] for row in t.rows]
        if rows:
            tables.append({"headers": rows[0], "rows": rows[1:51], "source": "docx"})
            parts.extend(" | ".join(r) for r in rows)
    return {"text": "\n".join(parts), "page_count": 1, "parser": "python-docx", "tables": tables}


def _parse_xlsx(content: bytes) -> dict:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    tables, parts = [], []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                rows.append(cells)
            if len(rows) > 51:
                break
        if rows:
            tables.append({"headers": rows[0], "rows": rows[1:51], "source": f"sheet:{ws.title}"})
            parts.append(f"[{ws.title}]")
            parts.extend(" | ".join(r) for r in rows)
    return {"text": "\n".join(parts), "page_count": len(wb.worksheets), "parser": "openpyxl", "tables": tables}


def _parse_image(content: bytes, filename: str) -> dict:
    from PIL import Image
    img = Image.open(io.BytesIO(content))
    info = {"format": img.format, "mode": img.mode, "width": img.width, "height": img.height,
            "megapixels": round(img.width * img.height / 1_000_000, 2),
            "orientation": "landscape" if img.width >= img.height else "portrait"}
    exif = getattr(img, "_getexif", lambda: None)()
    if exif:
        info["has_exif"] = True
    out = {"text": "", "page_count": 1, "parser": "pillow", "image_info": info}
    if ocr_available():
        import pytesseract
        out["text"] = (pytesseract.image_to_string(img) or "").strip()
        out["ocr_used"] = True
        info["ocr_chars"] = len(out["text"])
    else:
        out["needs_ocr"] = True
    return out


def extract_from_bytes(filename: str, content: bytes) -> dict:
    """Route a file to the right parser. Returns {text, page_count, parser,
    tables?, image_info?, ocr_used?, needs_ocr?} — never raises for an
    unsupported type; reports it instead."""
    name = (filename or "file").lower()
    ext = "." + name.rsplit(".", 1)[-1] if "." in name else ""
    if ext == ".pdf":
        return _parse_pdf(content)
    if ext == ".docx":
        return _parse_docx(content)
    if ext == ".xlsx":
        return _parse_xlsx(content)
    if ext in IMAGE_EXTS:
        return _parse_image(content, filename)
    # everything else: treat as text
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = content.decode("latin-1")
        except Exception:
            return {"text": "", "page_count": 0, "parser": "none",
                    "error": f"Unsupported file type '{ext or 'unknown'}'"}
    return {"text": text, "page_count": 1, "parser": "text"}


# ---------- classification ----------
DOC_SIGNALS: dict[str, list[str]] = {
    "invoice": ["invoice", "invoice no", "invoice number", "bill to", "gstin", "gst", "subtotal",
                "amount due", "total due", "payment due", "tax invoice", "hsn", "balance due"],
    "receipt": ["receipt", "payment received", "paid on", "transaction id", "change due", "cashier"],
    "contract": ["agreement", "party", "parties", "whereas", "hereinafter", "termination",
                 "governing law", "indemnify", "confidentiality", "effective date", "witness whereof"],
    "identity": ["passport", "aadhaar", "permanent account number", "pan card", "date of birth",
                 "government of india", "driving licence", "driver license", "nationality", "voter id"],
    "resume": ["curriculum vitae", "resume", "work experience", "professional experience", "education",
               "skills", "career objective", "certifications", "references available"],
    "report": ["executive summary", "quarterly report", "annual report", "analysis", "findings",
               "conclusion", "methodology", "appendix"],
    "letter": ["dear sir", "dear madam", "yours sincerely", "yours faithfully", "regards", "to whom it may concern"],
}


def classify_document(text: str, filename: str = "") -> tuple[str, float, dict]:
    hay = f"{filename}\n{text}".lower()
    scores = {t: sum(1 for kw in kws if kw in hay) for t, kws in DOC_SIGNALS.items()}
    best = max(scores, key=lambda t: scores[t])
    total = sum(scores.values())
    if scores[best] == 0:
        return "other", 0.0, scores
    confidence = round(scores[best] / max(total, 1), 3)
    return best, confidence, scores


# ---------- data extraction ----------
RE_EMAIL = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
RE_PHONE = re.compile(r"(?:\+?\d{1,3}[\s-]?)?(?:\(?\d{3,5}\)?[\s-]?)\d{3}[\s-]?\d{3,4}")
RE_DATE = re.compile(r"\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2}|"
                     r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4})\b",
                     re.IGNORECASE)
RE_AMOUNT = re.compile(r"(?:₹|Rs\.?|INR|USD|\$|EUR|€)\s?([\d,]+(?:\.\d{1,2})?)")
RE_GSTIN = re.compile(r"\b\d{2}[A-Z]{5}\d{4}[A-Z]\d[A-Z0-9]{2}\b")
RE_PAN = re.compile(r"\b[A-Z]{5}\d{4}[A-Z]\b")
RE_AADHAAR = re.compile(r"\b\d{4}\s\d{4}\s\d{4}\b")
RE_PASSPORT = re.compile(r"\b[A-Z]\d{7}\b")
RE_URL = re.compile(r"https?://[^\s)>\"]+")

SKILL_WORDS = ["python", "java", "javascript", "typescript", "react", "sql", "excel", "sales",
               "marketing", "negotiation", "crm", "communication", "leadership", "aws", "docker",
               "fastapi", "django", "node", "management", "accounting", "photoshop", "seo"]


def _mask(value: str) -> str:
    digits = re.sub(r"\s", "", value)
    return ("*" * max(0, len(digits) - 4)) + digits[-4:]


def _amounts(text: str) -> list[float]:
    out = []
    for m in RE_AMOUNT.finditer(text):
        try:
            out.append(float(m.group(1).replace(",", "")))
        except ValueError:
            pass
    return out


def _labeled(text: str, labels: list[str]) -> str | None:
    for label in labels:
        m = re.search(rf"{label}\s*[:#-]?\s*([^\n]{{1,80}})", text, re.IGNORECASE)
        if m:
            return m.group(1).strip().strip(":#- ") or None
    return None


def extract_common(text: str) -> dict:
    return {
        "emails": sorted(set(RE_EMAIL.findall(text)))[:10],
        "phones": sorted({p.strip() for p in RE_PHONE.findall(text) if sum(c.isdigit() for c in p) >= 9})[:10],
        "dates": [m.group(0) for m in RE_DATE.finditer(text)][:15],
        "amounts": _amounts(text)[:15],
        "urls": sorted(set(RE_URL.findall(text)))[:10],
    }


def extract_invoice(text: str) -> dict:
    amounts = _amounts(text)
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    total = None
    m = re.search(r"(?<![a-z])(?:grand\s+)?total(?:\s+due|\s+amount)?\s*[:\-]?\s*(?:₹|Rs\.?|INR|USD|\$|EUR|€)?\s?([\d,]+(?:\.\d{1,2})?)",
                  text, re.IGNORECASE)
    if m:
        try:
            total = float(m.group(1).replace(",", ""))
        except ValueError:
            pass
    tax = None
    m = re.search(r"(?:gst|tax|cgst|sgst|igst|vat)\s*(?:@?\s*\d{1,2}%?)?\s*[:\-]?\s*(?:₹|Rs\.?|INR|USD|\$)?\s?([\d,]+(?:\.\d{1,2})?)",
                  text, re.IGNORECASE)
    if m:
        try:
            tax = float(m.group(1).replace(",", ""))
        except ValueError:
            pass
    return {
        "invoice_number": _labeled(text, ["invoice\\s*(?:no|number|#)", "inv\\s*(?:no|#)", "bill\\s*(?:no|number)"]),
        "invoice_date": _labeled(text, ["invoice\\s*date", "date\\s*of\\s*issue", "dated?"]),
        "due_date": _labeled(text, ["due\\s*date", "payment\\s*due"]),
        "vendor": lines[0][:80] if lines else None,
        "bill_to": _labeled(text, ["bill\\s*to", "billed\\s*to", "customer"]),
        "gstin": (RE_GSTIN.findall(text) or [None])[0],
        "total": total if total is not None else (max(amounts) if amounts else None),
        "tax": tax,
        "currency": ("INR" if re.search(r"₹|Rs\.?|INR", text) else "USD" if "$" in text or "USD" in text else None),
    }


def extract_contract(text: str) -> dict:
    parties = None
    m = re.search(r"between\s+(.{3,80}?)\s+(?:and)\s+(.{3,80}?)(?:[\.,;\n]|$)", text, re.IGNORECASE)
    if m:
        parties = [m.group(1).strip().strip('",'), m.group(2).strip().strip('",')]
    notice = None
    m = re.search(r"(\d{1,3})\s*(?:calendar\s+|business\s+)?days[’']?\s*(?:written\s+)?notice", text, re.IGNORECASE)
    if m:
        notice = int(m.group(1))
    term = None
    m = re.search(r"(?:term|period)\s+of\s+(\d{1,3})\s*(months?|years?)", text, re.IGNORECASE)
    if m:
        term = f"{m.group(1)} {m.group(2)}"
    return {
        "parties": parties,
        "effective_date": _labeled(text, ["effective\\s*(?:as\\s*of|date)", "commencement\\s*date", "dated"]),
        "term": term,
        "termination_notice_days": notice,
        "auto_renewal": bool(re.search(r"auto[\s-]?renew", text, re.IGNORECASE)),
        "governing_law": _labeled(text, ["governing\\s*law", "governed\\s*by\\s*the\\s*laws\\s*of"]),
        "payment_terms": _labeled(text, ["payment\\s*terms", "net\\s*(?=\\d)"]),
        "confidentiality": bool(re.search(r"confidential", text, re.IGNORECASE)),
    }


def extract_identity(text: str) -> dict:
    """Identity document parsing. Numbers are stored MASKED (last 4 only)."""
    out: dict = {"id_numbers": []}
    for kind, rx in (("pan", RE_PAN), ("aadhaar", RE_AADHAAR), ("passport", RE_PASSPORT)):
        for v in rx.findall(text)[:3]:
            out["id_numbers"].append({"type": kind, "number_masked": _mask(v)})
    hay = text.lower()
    out["document_kind"] = ("aadhaar" if "aadhaar" in hay else
                            "pan" if "permanent account number" in hay or "pan" in hay else
                            "passport" if "passport" in hay else
                            "driving_licence" if "licence" in hay or "license" in hay else "unknown")
    out["name"] = _labeled(text, ["name"])
    out["date_of_birth"] = _labeled(text, ["date\\s*of\\s*birth", "dob"])
    out["nationality"] = _labeled(text, ["nationality"])
    return out


def extract_resume(text: str) -> dict:
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    hay = text.lower()
    years = None
    m = re.search(r"(\d{1,2})\+?\s*years?\s+(?:of\s+)?experience", hay)
    if m:
        years = int(m.group(1))
    education = [ln for ln in lines if re.search(
        r"\b(b\.?tech|m\.?tech|mba|b\.?sc|m\.?sc|b\.?com|m\.?com|bachelor|master|ph\.?d|diploma)\b",
        ln, re.IGNORECASE)][:5]
    return {
        "name": lines[0][:60] if lines else None,
        "email": (RE_EMAIL.findall(text) or [None])[0],
        "phone": (RE_PHONE.findall(text) or [None])[0],
        "skills": [s for s in SKILL_WORDS if s in hay][:15],
        "education": education,
        "experience_years": years,
    }


def extract_fields(doc_type: str, text: str) -> dict:
    data = {"common": extract_common(text)}
    if doc_type in ("invoice", "receipt"):
        data["invoice"] = extract_invoice(text)
    if doc_type == "contract":
        data["contract"] = extract_contract(text)
    if doc_type == "identity":
        data["identity"] = extract_identity(text)
    if doc_type == "resume":
        data["resume"] = extract_resume(text)
    return data


# ---------- table extraction (from plain text) ----------
def extract_tables_from_text(text: str) -> list[dict]:
    tables: list[dict] = []
    block: list[list[str]] = []

    def flush():
        nonlocal block
        if len(block) >= 2:
            width = max(len(r) for r in block)
            rows = [r + [""] * (width - len(r)) for r in block]
            tables.append({"headers": rows[0], "rows": rows[1:51], "source": "text"})
        block = []

    for line in text.splitlines():
        cells: list[str] | None = None
        if "|" in line and line.count("|") >= 2:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(re.fullmatch(r"[-: ]*", c) for c in cells):  # markdown separator row
                continue
        elif "\t" in line:
            cells = [c.strip() for c in line.split("\t")]
        elif re.search(r"\S\s{3,}\S", line):
            cells = [c.strip() for c in re.split(r"\s{3,}", line.strip())]
        if cells and len([c for c in cells if c]) >= 2:
            block.append(cells)
        else:
            flush()
    flush()
    return tables[:10]


class DocumentIntelligenceService:
    def __init__(self, db: AsyncSession):
        self.db = db
        self.audit = AuditService(db)

    def _require_manager(self, actor: User):
        if actor.role not in MANAGER_ROLES:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN,
                                detail="Manager or admin role required")

    async def _get(self, actor: User, doc_id: uuid.UUID) -> DIDocument:
        d = (await self.db.execute(select(DIDocument).filter(
            DIDocument.id == doc_id, DIDocument.organization_id == actor.organization_id,
            DIDocument.is_deleted == False))).scalars().first()
        if not d:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
        if actor.role not in MANAGER_ROLES and d.uploaded_by != actor.id:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN,
                                detail="You can only access documents you uploaded")
        return d

    # ---------- the pipeline ----------
    def _analyze(self, doc: DIDocument, text: str, parsed: dict):
        text = (text or "")[:TEXT_CAP]
        doc.text_content = text
        doc.page_count = int(parsed.get("page_count") or 0)
        doc.ocr_used = bool(parsed.get("ocr_used"))
        doc.image_info = parsed.get("image_info") or {}
        doc_type, confidence, signals = classify_document(text, doc.filename)
        doc.doc_type = doc_type
        doc.classification_confidence = confidence
        doc.classification_signals = signals
        doc.extraction = extract_fields(doc_type, text) if text else {}
        tables = list(parsed.get("tables") or [])
        if text:
            tables.extend(extract_tables_from_text(text))
        doc.tables = tables[:10]
        doc.embedding = embed_text(text[:20000]) if text else []
        doc.embedding_model = EMBED_MODEL
        if parsed.get("error"):
            doc.status = "failed"
            doc.error = parsed["error"]
        elif parsed.get("needs_ocr") and not text:
            doc.status = "needs_ocr"
            doc.error = ("No text layer found and OCR is not available on this server"
                         if not ocr_available() else
                         "Scanned PDF: page rasterization for OCR is not available")
        else:
            doc.status = "processed"
            doc.error = None
        doc.processed_at = _now()

    async def process_bytes(self, actor: User, filename: str, content: bytes, *,
                            content_type: str | None = None, source: str = "upload",
                            context_type: str | None = None,
                            context_id: uuid.UUID | None = None) -> dict:
        if not content:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Empty file")
        if len(content) > 15 * 1024 * 1024:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="File too large (max 15 MB)")
        doc = DIDocument(organization_id=actor.organization_id, filename=filename or "file",
                         content_type=content_type, size_bytes=len(content), source=source,
                         context_type=context_type, context_id=context_id, uploaded_by=actor.id)
        try:
            parsed = extract_from_bytes(filename, content)
        except Exception as e:  # a corrupt file must not 500 the endpoint
            parsed = {"text": "", "page_count": 0, "parser": "none", "error": f"Parse error: {e}"}
        self._analyze(doc, parsed.get("text") or "", parsed)
        self.db.add(doc)
        await self.db.flush()
        await self.audit.log_event(actor.organization_id, actor_user_id=actor.id,
                                   action="DOCUMENT_PROCESSED", resource_type="document_intelligence",
                                   resource_id=str(doc.id),
                                   action_metadata={"filename": doc.filename, "doc_type": doc.doc_type,
                                                    "status": doc.status, "ocr_used": doc.ocr_used})
        await self.db.commit()
        await self.db.refresh(doc)
        return self._doc_dict(doc)

    async def process_text(self, actor: User, text: str, *, filename: str = "pasted.txt",
                           context_type: str | None = None,
                           context_id: uuid.UUID | None = None) -> dict:
        if not (text or "").strip():
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="text is required")
        return await self.process_bytes(actor, filename, text.encode("utf-8"),
                                        content_type="text/plain", source="text",
                                        context_type=context_type, context_id=context_id)

    async def reprocess(self, actor: User, doc_id: uuid.UUID) -> dict:
        """Re-run classification/extraction/tables/embedding over the stored
        extracted text (original bytes are not retained)."""
        d = await self._get(actor, doc_id)
        self._analyze(d, d.text_content or "", {"page_count": d.page_count,
                                                "ocr_used": d.ocr_used,
                                                "image_info": d.image_info})
        await self.db.commit()
        await self.db.refresh(d)
        return self._doc_dict(d)

    async def summarize(self, actor: User, doc_id: uuid.UUID, length: int = 5) -> dict:
        d = await self._get(actor, doc_id)
        if not (d.text_content or "").strip():
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST,
                                detail="Document has no extracted text to summarize")
        from app.services.ai_gateway_service import AIGatewayService
        out = await AIGatewayService(self.db).summarize_text(actor, d.text_content, length=length)
        d.summary = out.get("text")
        await self.db.commit()
        return {"id": str(d.id), "summary": d.summary, "model": out.get("model"),
                "provider": out.get("provider"), "cached": out.get("cached", False)}

    # ---------- listing / search ----------
    def _visible_filter(self, query, actor: User):
        if actor.role not in MANAGER_ROLES:
            query = query.filter(DIDocument.uploaded_by == actor.id)
        return query

    async def list_documents(self, actor: User, *, doc_type: str | None = None,
                             status_f: str | None = None, q: str | None = None,
                             context_type: str | None = None,
                             context_id: uuid.UUID | None = None,
                             limit: int = 50, offset: int = 0) -> dict:
        query = select(DIDocument).filter(
            DIDocument.organization_id == actor.organization_id,
            DIDocument.is_deleted == False)
        query = self._visible_filter(query, actor)
        if doc_type:
            query = query.filter(DIDocument.doc_type == doc_type)
        if status_f:
            query = query.filter(DIDocument.status == status_f)
        if context_type:
            query = query.filter(DIDocument.context_type == context_type)
        if context_id:
            query = query.filter(DIDocument.context_id == context_id)
        if q:
            like = f"%{q}%"
            query = query.filter(DIDocument.filename.ilike(like) | DIDocument.text_content.ilike(like))
        docs = (await self.db.execute(query.order_by(DIDocument.created_at.desc()))).scalars().all()
        total = len(docs)
        return {"total": total,
                "items": [self._doc_dict(d, include_text=False) for d in docs[offset:offset + limit]]}

    async def get_document(self, actor: User, doc_id: uuid.UUID) -> dict:
        return self._doc_dict(await self._get(actor, doc_id))

    async def delete_document(self, actor: User, doc_id: uuid.UUID) -> dict:
        d = await self._get(actor, doc_id)
        d.is_deleted = True
        d.deleted_at = _now()
        await self.audit.log_event(actor.organization_id, actor_user_id=actor.id,
                                   action="DOCUMENT_DELETED", resource_type="document_intelligence",
                                   resource_id=str(d.id), action_metadata={"filename": d.filename})
        await self.db.commit()
        return {"deleted": True}

    async def search(self, actor: User, query_text: str, *, doc_type: str | None = None,
                     limit: int = 10) -> dict:
        qvec = embed_text(query_text)
        qtokens = set(_tokenize(query_text))
        q = select(DIDocument).filter(
            DIDocument.organization_id == actor.organization_id,
            DIDocument.is_deleted == False)
        q = self._visible_filter(q, actor)
        if doc_type:
            q = q.filter(DIDocument.doc_type == doc_type)
        docs = (await self.db.execute(q)).scalars().all()
        scored = []
        for d in docs:
            sim = cosine(qvec, d.embedding or [])
            dtokens = set(_tokenize(f"{d.filename} {(d.text_content or '')[:20000]}"))
            overlap = len(qtokens & dtokens) / max(1, len(qtokens))
            if overlap == 0 and sim < 0.25:
                continue
            score = 0.65 * sim + 0.35 * overlap
            excerpt = ""
            for tok in qtokens:
                idx = (d.text_content or "").lower().find(tok)
                if idx >= 0:
                    excerpt = (d.text_content or "")[max(0, idx - 60):idx + 180]
                    break
            scored.append({"id": str(d.id), "filename": d.filename, "doc_type": d.doc_type,
                           "score": round(score, 4), "excerpt": excerpt or (d.text_content or "")[:180]})
        scored.sort(key=lambda s: s["score"], reverse=True)
        return {"query": query_text, "results": scored[:limit], "count": min(len(scored), limit),
                "embedding_model": EMBED_MODEL, "search_type": "semantic_hybrid"}

    # ---------- dashboard / export ----------
    async def dashboard(self, actor: User) -> dict:
        q = select(DIDocument).filter(
            DIDocument.organization_id == actor.organization_id,
            DIDocument.is_deleted == False)
        q = self._visible_filter(q, actor)
        docs = (await self.db.execute(q.order_by(DIDocument.created_at.desc()))).scalars().all()
        by_type: dict[str, int] = {}
        by_status: dict[str, int] = {}
        ocr_used = pages = with_tables = with_extraction = 0
        for d in docs:
            by_type[d.doc_type] = by_type.get(d.doc_type, 0) + 1
            by_status[d.status] = by_status.get(d.status, 0) + 1
            ocr_used += 1 if d.ocr_used else 0
            pages += d.page_count or 0
            with_tables += 1 if d.tables else 0
            typed = {k: v for k, v in (d.extraction or {}).items() if k != "common"}
            with_extraction += 1 if typed else 0
        return {
            "totals": {"documents": len(docs), "by_type": by_type, "by_status": by_status,
                       "pages": pages, "ocr_used": ocr_used, "with_tables": with_tables,
                       "with_structured_extraction": with_extraction},
            "recent": [self._doc_dict(d, include_text=False) for d in docs[:8]],
            "capabilities": capabilities(),
        }

    async def export_csv(self, actor: User) -> str:
        self._require_manager(actor)
        docs = (await self.db.execute(select(DIDocument).filter(
            DIDocument.organization_id == actor.organization_id,
            DIDocument.is_deleted == False).order_by(DIDocument.created_at))).scalars().all()
        buf = io.StringIO()
        w = csv.writer(buf)
        w.writerow(["id", "filename", "doc_type", "confidence", "status", "source", "pages",
                    "ocr_used", "tables", "size_bytes", "uploaded_by", "created_at"])
        for d in docs:
            w.writerow([str(d.id), d.filename, d.doc_type, d.classification_confidence, d.status,
                        d.source, d.page_count, d.ocr_used, len(d.tables or []), d.size_bytes,
                        str(d.uploaded_by) if d.uploaded_by else "",
                        _aware(d.created_at).isoformat() if d.created_at else ""])
        await self.audit.log_event(actor.organization_id, actor_user_id=actor.id,
                                   action="DOCUMENT_EXPORTED", resource_type="document_intelligence",
                                   action_metadata={"rows": len(docs)})
        await self.db.commit()
        return buf.getvalue()

    # ---------- helpers ----------
    def _doc_dict(self, d: DIDocument, include_text: bool = True) -> dict:
        out = {"id": str(d.id), "filename": d.filename, "content_type": d.content_type,
               "size_bytes": d.size_bytes, "source": d.source,
               "context_type": d.context_type,
               "context_id": str(d.context_id) if d.context_id else None,
               "status": d.status, "error": d.error, "page_count": d.page_count,
               "ocr_used": d.ocr_used, "doc_type": d.doc_type,
               "classification_confidence": d.classification_confidence,
               "extraction": d.extraction or {}, "tables": d.tables or [],
               "image_info": d.image_info or {}, "summary": d.summary,
               "embedding_model": d.embedding_model,
               "uploaded_by": str(d.uploaded_by) if d.uploaded_by else None,
               "processed_at": _aware(d.processed_at).isoformat() if d.processed_at else None,
               "created_at": _aware(d.created_at).isoformat() if d.created_at else None}
        if include_text:
            out["text_content"] = d.text_content
        return out
